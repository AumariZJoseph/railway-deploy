import os
import fitz  # PyMuPDF for PDF processing
import pandas as pd
from pathlib import Path
from docx import Document as DocxDocument
import hashlib
import re

def extract_text_with_metadata(file_path, filename):
    """Extract text with better structure and metadata"""
    text, error = process_file(file_path, filename)
    if error:
        return None, error
    
    enhanced_text = f"""
DOCUMENT: {filename}
CONTENT:
{text}
"""
    return enhanced_text.strip(), None

# ---------------- PDF HANDLING ----------------
def process_pdf_file(file_path):
    """Enhanced PDF processing with accurate page labeling"""
    try:
        full_text = ""
        doc = fitz.open(file_path)
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            
            if text.strip():
                full_text += f"Page {page_num + 1}:\n{text}\n\n"
        
        doc.close()
        
        if not full_text or len(full_text.strip()) < 50:
            return None, "Document appears to be empty or contains no extractable text"
        return full_text, None
    except Exception as e:
        error_msg = f"Error processing PDF {file_path}: {str(e)}"
        print(error_msg)
        return None, error_msg

# ---------------- WORD HANDLING ----------------
def process_word_document(file_path):
    """Enhanced Word document processing with structure preservation"""
    try:
        doc = DocxDocument(file_path)
        full_text = ""
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Headings
            if paragraph.style and paragraph.style.name.startswith('Heading'):
                try:
                    level = int(paragraph.style.name.split()[-1]) if 'Heading' in paragraph.style.name else 1
                except:
                    level = 1
                header_prefix = "#" * min(level, 3) + " "
                full_text += f"\n{header_prefix}{text}\n"
                continue
            
            # Regular text
            full_text += text + "\n\n"
        
        # Tables
        for table in doc.tables:
            full_text += "\n[Table Start]\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = " ".join(par.text.strip() for par in cell.paragraphs)
                    row_text.append(cell_text)
                full_text += " | ".join(row_text) + "\n"
            full_text += "[Table End]\n\n"
        
        if not full_text or len(full_text.strip()) < 50:
            return None, "Document appears to be empty or contains minimal text"
        return full_text, None
    except Exception as e:
        error_msg = f"Error processing Word document {file_path}: {str(e)}"
        print(error_msg)
        return None, error_msg


# ---------------- TXT HANDLING ----------------
def process_text_file(file_path):
    """Process plain text files with structure preservation"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
            return content, None
    except Exception as e:
        error_msg = f"Error processing text file {file_path}: {str(e)}"
        print(error_msg)
        return None, error_msg

# ---------------- FILE HANDLER ----------------
def process_file(file_path, filename):
    """Process a file with enhanced type handling"""
    file_ext = Path(filename).suffix.lower()
    if file_ext == '.pdf':
        return process_pdf_file(file_path)
    elif file_ext == '.docx':
        return process_word_document(file_path)
    elif file_ext == '.txt':
        return process_text_file(file_path)
    else:
        return None, "Unsupported file type"

# ---------------- METADATA ----------------
def get_file_metadata(file_path):
    """Get metadata for a file"""
    return {
        "hash": calculate_file_hash(file_path),
        "size": os.path.getsize(file_path),
        "modified": os.path.getmtime(file_path)
    }

def calculate_file_hash(file_path):
    """Calculate MD5 hash of a file"""
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()